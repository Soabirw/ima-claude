"""Inspect Word doc styles and paragraph structure."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

if len(sys.argv) < 2:
    print("Usage: python3 inspect_docx.py <input.docx>")
    sys.exit(1)

doc = Document(sys.argv[1])

for i, p in enumerate(doc.paragraphs[:80]):
    if not p.text.strip():
        continue
    s = p.style.name if p.style else 'None'
    bold_runs = any(r.bold for r in p.runs if r.text.strip())
    all_bold = all(r.bold for r in p.runs if r.text.strip()) if p.runs else False
    bstr = 'ALL' if all_bold else ('mix' if bold_runs else '   ')
    preview = p.text.strip()[:90]
    print(f"{i:3d} [{s:20s}] [{bstr}] {preview}")

print()
styles = sorted(set(p.style.name for p in doc.paragraphs if p.style and p.text.strip()))
print(f"All styles: {styles}")

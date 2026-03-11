"""
Generate a styled IMA Cancer Care companion guide as HTML from a Word doc.

Styling matches the Canva "Cancer Drug Resistance Guide March 2026" design.
Values extracted from Canva design data (176 font-size declarations, 94 elements).

Usage:
    /c/Python313/python.exe generate_html.py <path_to_docx> [--out <output.html>]

Python path: /c/Python313/python.exe
Required: pip install python-docx
"""

import sys
import io
import base64
from pathlib import Path
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent))
from extract_docx import extract_document
from docx import Document


# ── Brand CSS — values confirmed from Canva design data ──────────────────────
# Canva internal units × 0.75 = points
# Cover title 1: 120.722 × 0.75 = 90.5pt
# Cover title 2: 89.293 × 0.75 = 67pt
# Section heading: 20.053 × 0.75 = 15pt
# Sub heading: 17.333 × 0.75 = 13pt
# Body text: 16.0 × 0.75 = 12pt
# Footer: 13.333 × 0.75 = 10pt
# Authors: 25.333 × 0.75 = 19pt
# Subtitle: 33.333 × 0.75 = 25pt

CSS = """
@import url('https://fonts.googleapis.com/css2?family=Lato:ital,wght@0,400;0,700;0,900;1,400;1,700&display=swap');

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

:root {
    --navy: #00066F;
    --teal: #00B8B8;
    --gold: #FFCC00;
    --black: #000000;
    --dark-gray: #333333;
    --light-gray: #D9D9D9;
    --white: #FFFFFF;
    --page-width: 8.5in;
    --page-height: 11in;
}

@page {
    size: letter;
    margin: 0.5in;
}

html {
    background: #525659;
}

body {
    font-family: 'Lato', sans-serif;
    font-size: 12pt;
    line-height: 1.4;
    color: var(--black);
    max-width: var(--page-width);
    margin: 0 auto;
    padding: 0;
    letter-spacing: 0;  /* Canva tracking=0 for body text */
    background: transparent;
}

.page {
    width: var(--page-width);
    min-height: var(--page-height);
    background: var(--white);
    padding: 0.5in 1.05in 0.5in 0.69in;
    margin: 20px auto;
    box-shadow: 0 2px 8px rgba(0,0,0,0.3);
    position: relative;
    overflow: hidden;
}

.page.cover-page {
    padding: 0;
}

@media print {
    html { background: white; }
    .page {
        margin: 0;
        box-shadow: none;
        page-break-after: always;
    }
}

/* ── Cover Page — absolute positioning from Canva coords ── */
/* Canva page: 8.5 × 11in. All positions in inches.        */
.cover {
    width: var(--page-width);
    height: var(--page-height);
    position: relative;
    background: var(--white);
    overflow: hidden;
}

/* Logo: 3.4×0.82in at X:2.55 Y:0.44 */
.cover-logo {
    position: absolute;
    left: 2.55in;
    top: 0.44in;
    width: 3.4in;
    height: 0.82in;
}
.cover-logo img {
    width: 100%;
    height: 100%;
    object-fit: contain;
}

/* CANCER title: 5.88×1.5in at X:1.33 Y:2.29 — behind blue box */
.cover-title-1-wrap {
    position: absolute;
    left: 1.33in;
    top: 2.15in;
    width: 5.88in;
    height: 1.5in;
    z-index: 1;
    display: flex;
    align-items: flex-end;
    justify-content: center;
}
.cover-title-1 {
    font-size: 90.5pt;
    font-weight: 700;
    line-height: 1;
    text-transform: uppercase;
    letter-spacing: 0.03em;
    color: #494949; /* Gravel */
    position: relative;
    top: 3pt;
    display: inline-block;
}
.cover-title-1::before {
    content: attr(data-text);
    position: absolute;
    left: 3pt;
    top: 1.4pt;
    color: #919396; /* Calm Stone */
    z-index: -1;
}

/* Navy blue box: 8.51×3.86in at X:-0.01 Y:3.56 */
.cover-hero {
    position: absolute;
    left: -0.01in;
    top: 3.56in;
    width: 8.51in;
    height: 3.86in;
    background: var(--navy);
    z-index: 2;
}

/* RESISTANCE: 6.92×1.3in at X:0.98 Y:3.37 — starts above blue box */
.cover-title-2 {
    position: absolute;
    left: 0.98in;
    top: 3.37in;
    width: 6.92in;
    height: 1.3in;
    z-index: 3;
    font-size: 78.2pt;
    font-weight: 700;
    line-height: 1;
    text-transform: uppercase;
    letter-spacing: 0.03em;
    color: var(--white);
    text-align: center;
    display: flex;
    align-items: center;
    justify-content: center;
}

/* Subtitle: 8.09×0.41in at X:0.22 Y:4.89 */
.cover-subtitle {
    position: absolute;
    left: 0.22in;
    top: 4.89in;
    width: 8.09in;
    height: 0.41in;
    z-index: 3;
    font-size: 25pt;
    font-weight: 400;
    color: var(--white);
    text-align: center;
    line-height: 1.3;
    letter-spacing: 0.02em;
}

/* Authors: 8.5×0.78in at X:-0.01 Y:6.38 */
.cover-authors {
    position: absolute;
    left: -0.01in;
    top: 6.38in;
    width: 8.5in;
    height: 0.78in;
    z-index: 3;
    font-size: 20pt;
    font-weight: 700;
    color: var(--white);
    text-align: center;
    line-height: 1.6;
}

/* Thumbnail: 1.46×1.9in at X:0.39 Y:8.06 rotate:-13.8deg */
img.cover-info {
    position: absolute;
    left: 0.39in;
    top: 8.06in;
    width: 1.46in;
    height: 1.9in;
    z-index: 2;
    object-fit: cover;
    box-shadow: 0 2px 8px rgba(0,0,0,0.15);
    transform: rotate(-13.8deg);
}

/* Disclaimer: 5.6×1.11in at X:2.49 Y:8.6 */
.cover-info-text {
    position: absolute;
    left: 2.49in;
    top: 8.6in;
    width: 5.6in;
    height: 1.11in;
    z-index: 2;
    font-size: 12pt;
    color: var(--dark-gray);
    line-height: 1.4;
    text-align: left;
}
.cover-info-text a {
    color: var(--navy);
}

/* Date: 8.49×0.2in at X:0 Y:10.69 */
.cover-date {
    position: absolute;
    left: 0in;
    top: 10.69in;
    width: 8.49in;
    height: 0.2in;
    z-index: 2;
    font-size: 12pt;
    font-weight: 700;
    font-style: italic;
    color: var(--dark-gray);
    text-align: center;
}

/* ── Content Headings ───────────────────────────────────── */
h2, .section-heading {
    font-size: 15pt;
    font-weight: 700;
    color: var(--navy);
    text-align: center;
    margin: 20pt 0 8pt;
    line-height: 1.4;
    letter-spacing: 0;
    break-before: page;
    page-break-before: always;
}

/* First h2 after cover — already on new page, don't double-break */
h2.no-page-break {
    break-before: auto;
    page-break-before: auto;
}

h3, .sub-heading {
    font-size: 13pt;
    font-weight: 700;
    color: var(--navy);
    text-align: left;
    margin: 16pt 0 8pt;
    line-height: 1.4;
    letter-spacing: 0;
}

/* ── Body Text ──────────────────────────────────────────── */
p, .body-text {
    font-size: 12pt;
    font-weight: 400;
    color: var(--black);
    line-height: 1.4;
    margin: 0 0 16.8pt;
    text-align: left;
}

/* Inline bold navy — sub-sub headings within body */
.bold-navy {
    font-weight: 700;
    color: var(--navy);
}

/* Roman numeral sub-section labels (I., II., III.) */
p.roman-label {
    margin-top: 18pt;
    margin-bottom: 0;
}

/* Remove gap between lead-in text and its list */
p.lead-in {
    margin-bottom: 0;
}

p + ul, p + ol {
    margin-top: 0;
}

ul, ol {
    margin: 2pt 0 16.8pt 18pt;
    padding: 0;
}

li {
    font-size: 12pt;
    color: var(--black);
    line-height: 1.4;
    margin-bottom: 0;
}

li::marker {
    color: var(--navy);
}

ul ul {
    margin-left: 18pt;
    margin-top: 1pt;
}

/* ── Page Breaks ───────────────────────────────────────── */
.page-break {
    break-before: page;
    page-break-before: always;
    border: none;
    border-top: 1px solid var(--light-gray);
    margin: 24pt 0;
    height: 0;
}

@media print {
    .page-break {
        border: none;
        margin: 0;
    }
}

/* ── Warning Box: 5.37×1.16in at X:1.55 Y:8.07 (centered, flow) ── */
.warning-box {
    background: linear-gradient(150deg, #00066F 0%, #00B8B8 100%);
    color: var(--white);
    font-size: 15.5pt;
    font-weight: 700;
    text-align: center;
    padding: 18pt 24pt;
    margin: 140pt auto 0;
    max-width: 5.37in;
    line-height: 1.5;
    border-radius: 4px;
}

.warning-box .gold {
    color: var(--gold);
}

/* ── Disclaimer ─────────────────────────────────────────── */
.disclaimer {
    font-size: 12pt;
    color: var(--black);
    text-align: center;
    font-style: italic;
    margin: 6pt 0;
    padding: 8pt 16pt;
    border-left: 3px solid var(--navy);
    text-align: left;
}

/* ── Q&A ────────────────────────────────────────────────── */
.qa-question {
    font-size: 12pt;
    font-weight: 700;
    color: var(--navy);
    margin: 10pt 0 2pt;
}

.qa-answer {
    font-size: 12pt;
    color: var(--black);
    margin: 2pt 0 4pt;
    text-align: left;
}

/* ── Captions ───────────────────────────────────────────── */
.caption {
    font-size: 12pt;
    color: var(--black);
    margin: 16pt 0 4pt;
}

/* ── Figures ────────────────────────────────────────────── */
.figure {
    text-align: center;
    margin: 12pt 0;
    page-break-inside: avoid;
}

.figure img {
    max-width: 100%;
    height: auto;
}

.figure-caption {
    font-size: 12pt;
    color: var(--black);
    margin-top: 4pt;
    font-style: italic;
}

/* ── References ─────────────────────────────────────────── */
.ref-heading {
    font-size: 13pt;
    font-weight: 700;
    color: var(--navy);
    margin: 10pt 0 8pt;
}

.reference {
    font-size: 8pt;
    color: var(--dark-gray);
    line-height: 1.375;
    margin: 1pt 0;
    padding-left: 14pt;
    text-indent: -14pt;
}

/* ── Footer: 8.5×0.16in at X:-0.02 Y:10.65 ────────────── */
.page-footer {
    position: absolute;
    left: 0;
    bottom: 0.2in;
    width: 100%;
    padding: 0 0.75in;
    font-size: 10pt;
    text-align: center;
    letter-spacing: 0.05em;
}

.page-footer .footer-title {
    color: var(--black);
}

.page-footer .footer-date {
    color: var(--black);
}

/* ── Page Breaks ────────────────────────────────────────── */
.page-break {
    break-before: page;
    page-break-before: always;
    height: 0;
    margin: 0;
    padding: 0;
}

h2 {
    break-before: auto;
}

/* ── Print ──────────────────────────────────────────────── */
@media print {
    body { padding: 0; max-width: none; }
    .cover { margin: 0; min-height: auto; height: 100vh; }
    .page-break { break-before: page; page-break-before: always; }
    .no-break { break-inside: avoid; page-break-inside: avoid; }
    .figure { break-inside: avoid; page-break-inside: avoid; }
    h2, h3 { break-after: avoid; page-break-after: avoid; }
}
"""


def safe(text):
    return (text.replace("&", "&amp;")
               .replace("<", "&lt;")
               .replace(">", "&gt;"))


def runs_to_html(runs):
    if not runs:
        return ""
    out = []
    for r in runs:
        t = safe(r.get("text", ""))
        if not t:
            continue
        if r.get("bold") and r.get("italic"):
            out.append(f"<strong><em>{t}</em></strong>")
        elif r.get("bold"):
            out.append(f"<strong>{t}</strong>")
        elif r.get("italic"):
            out.append(f"<em>{t}</em>")
        else:
            out.append(t)
    return "".join(out)


def para_html(entry):
    markup = runs_to_html(entry.get("runs", []))
    return markup if markup else safe(entry["text"])


def extract_images(docx_path):
    """Extract images from docx, return dict of rId -> base64 data URI."""
    doc = Document(docx_path)
    images = {}
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            ct = rel.target_part.content_type
            b64 = base64.b64encode(rel.target_part.blob).decode('ascii')
            images[rel.rId] = f"data:{ct};base64,{b64}"
    return images


def find_image_positions(docx_path):
    """Find which paragraph indices contain images and their rIds."""
    doc = Document(docx_path)
    positions = {}  # para_index -> [rId, ...]
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            drawings = run._element.findall(qn('w:drawing'))
            for d in drawings:
                blips = d.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        positions.setdefault(i, []).append(embed)
    return positions


def extract_cover_meta(sections):
    """Extract cover page metadata from document sections."""
    title_parts = []
    subtitle = ""
    authors = ""
    disclaimer = ""
    date_str = ""
    got_title = False

    for entry in sections[:30]:
        t = entry["type"]
        text = entry["text"].strip()

        if t in ("h1", "heading_bold"):
            if not got_title:
                title_parts.append(text)
                if len(title_parts) >= 2:
                    got_title = True
            else:
                if text.startswith("and ") or len(text) > 40:
                    subtitle = text
                else:
                    break
        elif t == "h2":
            if not got_title and not title_parts:
                subtitle = text
            elif text.startswith("and "):
                subtitle = text
            elif "M.D." in text or "MD" in text:
                authors = (authors + " | " + text) if authors else text
            else:
                break
        elif t == "author":
            authors = (authors + " | " + text) if authors else text
        elif t == "disclaimer":
            disclaimer = text
        elif t == "date":
            date_str = text
        elif t == "body" and got_title:
            break

    if not authors:
        authors = "Paul E. Marik, MD, FCCM, FCCP  |  Justus R. Hope, MD"

    return title_parts, subtitle, authors, disclaimer, date_str


def load_cover_asset(filename):
    """Load a base64 asset file from the scripts directory."""
    asset_path = Path(__file__).parent / filename
    if asset_path.exists():
        return asset_path.read_text().strip()
    return ""


def generate_html(docx_path, out_path):
    print(f"Extracting: {docx_path}")
    data = extract_document(str(docx_path))

    print("Extracting images...")
    images = extract_images(str(docx_path))
    image_positions = find_image_positions(str(docx_path))
    print(f"  Found {len(images)} images at {len(image_positions)} positions")

    # Cover metadata
    title_parts, subtitle, authors, disclaimer, date_str = extract_cover_meta(
        data["sections"]
    )

    # Build HTML
    html = []
    html.append("<!DOCTYPE html>")
    html.append('<html lang="en">')
    html.append("<head>")
    html.append('<meta charset="UTF-8">')
    html.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    html.append(f"<title>{safe(' '.join(title_parts))}</title>")
    html.append(f"<style>{CSS}</style>")
    html.append("</head>")
    html.append("<body>")

    # ── Cover ──────────────────────────────────────────────
    logo_b64 = load_cover_asset("logo_b64.txt")
    thumb_b64 = load_cover_asset("thumb_b64.txt")

    html.append('<div class="page cover-page"><div class="cover">')

    # Logo — absolute positioned
    if logo_b64:
        html.append(f'<div class="cover-logo"><img src="{logo_b64}" alt="Independent Medical Alliance"></div>')

    # CANCER title — absolute positioned, overlaps into blue box
    if title_parts:
        html.append(f'<div class="cover-title-1-wrap"><div class="cover-title-1" data-text="{safe(title_parts[0])}">{safe(title_parts[0])}</div></div>')

    # Navy blue box background — absolute positioned
    html.append('<div class="cover-hero"></div>')

    # RESISTANCE — absolute positioned, starts above blue box
    for idx, part in enumerate(title_parts):
        if idx == 0:
            continue
        html.append(f'<div class="cover-title-2">{safe(part)}</div>')

    # Subtitle — absolute positioned
    if subtitle:
        html.append(f'<div class="cover-subtitle">{safe(subtitle)}</div>')

    # Authors — absolute positioned
    if authors:
        author_list = [a.strip() for a in authors.replace("|", "\n").split("\n") if a.strip()]
        html.append('<div class="cover-authors">')
        html.append("<br>".join(safe(a) for a in author_list))
        html.append('</div>')

    # Thumbnail — absolute positioned
    if thumb_b64:
        html.append(f'<img class="cover-info" src="{thumb_b64}" alt="Cancer Care Guide">')

    # Disclaimer — absolute positioned
    if disclaimer:
        html.append(f'<div class="cover-info-text">{safe(disclaimer)}</div>')

    # Date — absolute positioned
    if date_str:
        html.append(f'<div class="cover-date">{safe(date_str)}</div>')

    html.append("</div>")  # cover
    html.append("</div>")  # page cover-page

    # Start first content page
    html.append('<div class="page">')

    # ── Build image lookup: for each section entry, check if images fall
    #    between it and the next entry (covers image-only paragraphs skipped
    #    by extract_docx because they have no text) ────────────────────────
    def insert_pending_images(after_idx, before_idx, html_out):
        """Insert any images whose paragraph index is between after_idx and before_idx."""
        for img_idx in sorted(image_positions.keys()):
            if img_idx in rendered_images:
                continue
            if img_idx > after_idx and (before_idx is None or img_idx < before_idx):
                for rId in image_positions[img_idx]:
                    if rId in images:
                        html_out.append('<div class="figure">')
                        html_out.append(f'<img src="{images[rId]}" alt="Figure">')
                        html_out.append("</div>")
                rendered_images.add(img_idx)

    # ── Content ────────────────────────────────────────────
    sections = data["sections"]
    cover_titles = set(p.lower() for p in title_parts)
    h2_content_count = 0  # track h2s rendered in content (not cover)
    in_list = False
    rendered_images = set()
    num_counters = {}  # numId -> current count, for Word auto-numbering
    i = 0

    def next_is_bullet(idx):
        """Check if the next section entry after idx is a bullet."""
        return idx + 1 < len(sections) and sections[idx + 1]["type"] == "bullet"

    while i < len(sections):
        entry = sections[i]
        t = entry["type"]
        text = entry["text"].strip()
        para_idx = entry.get("index", -1)

        # Close open list if this isn't a bullet
        if t != "bullet" and in_list:
            html.append("</ul>")
            in_list = False

        # Insert any images that fall between the previous section entry
        # and this one (catches image-only paragraphs with no text)
        if i > 0:
            prev_idx = sections[i - 1].get("index", -1)
            insert_pending_images(prev_idx, para_idx, html)

        # Check if this paragraph itself has an image
        if para_idx in image_positions and para_idx not in rendered_images:
            for rId in image_positions[para_idx]:
                if rId in images:
                    html.append('<div class="figure">')
                    html.append(f'<img src="{images[rId]}" alt="Figure">')
                    html.append("</div>")
            rendered_images.add(para_idx)
            if not text:
                i += 1
                continue

        # Skip cover metadata already rendered (before page_break check
        # so skipped entries don't create blank pages)
        if t in ("h1", "author", "date", "disclaimer"):
            i += 1
            continue

        # Page breaks from Word document
        if t == "page_break":
            html.append('</div><div class="page">')
            i += 1
            continue
        # Honour w:br page breaks on content paragraphs (break before)
        # but not on warnings (handled separately in the warning block)
        if entry.get("page_break") and t != "warning":
            html.append('</div><div class="page">')

        if t == "heading_bold":
            # Skip cover elements
            if text.lower() in cover_titles:
                i += 1
                continue
            if text.startswith("and ") or text.startswith("Paul E.") or text.startswith("Justus"):
                i += 1
                continue
            # Roman numeral sub-labels get roman-label class
            import re as _re
            if _re.match(r'^[IVX]+[\.\s]', text):
                classes = ["roman-label"]
                if next_is_bullet(i):
                    classes.append("lead-in")
                html.append(f'<p class="{" ".join(classes)}"><strong>{safe(text)}</strong></p>')
                i += 1
                continue
            # Render as bold body text — actual section headings use Word
            # Heading styles (h2/h3), heading_bold are inline bold labels
            li = ' class="lead-in"' if next_is_bullet(i) else ""
            html.append(f"<p{li}><strong>{safe(text)}</strong></p>")
            i += 1
            continue

        if t == "h2":
            if text.startswith("and ") or "M.D." in text or ", MD" in text or text.endswith("MD"):
                i += 1
                continue
            h2_content_count += 1
            if h2_content_count > 1:
                html.append('</div><div class="page">')
            html.append(f"<h2>{para_html(entry)}</h2>")
            i += 1
            continue

        if t == "h3":
            prefix = ""
            if entry.get("numId"):
                nid = entry["numId"]
                num_counters[nid] = num_counters.get(nid, 0) + 1
                prefix = f"{num_counters[nid]}. "
            # Roman numeral sub-labels (I., II., III.) → bold body, not heading
            import re as _re
            if _re.match(r'^[IVX]+[\.\s]', text):
                classes = ["roman-label"]
                if next_is_bullet(i):
                    classes.append("lead-in")
                html.append(f'<p class="{" ".join(classes)}"><strong>{safe(text)}</strong></p>')
            else:
                html.append(f"<h3>{safe(prefix)}{para_html(entry)}</h3>")
            i += 1
            continue

        if t == "ref_heading":
            # Skip — references are rendered in the dedicated block below
            i += 1
            continue

        if t == "warning":
            # The warning box should only contain the core message.
            # Extra sentences may have been merged in from the docx.
            import re as _re
            # Split at end of "patients should not treat themselves." sentence
            m = _re.search(r'patients should not treat themselves\.?', text, flags=_re.IGNORECASE)
            if m:
                warning_text = text[:m.end()].strip()
                remainder = text[m.end():].strip()
            else:
                warning_text = text
                remainder = ""
            warning_html = safe(warning_text)
            warning_html = _re.sub(
                r'(patients should not treat themselves)',
                r'<span class="gold">\1</span>',
                warning_html,
                flags=_re.IGNORECASE
            )
            html.append(f'<div class="warning-box">{warning_html}</div>')
            # Honour the Word page break that sits on the warning paragraph
            if entry.get("page_break") or remainder:
                html.append('</div><div class="page">')
            if remainder:
                html.append(f"<p>{safe(remainder)}</p>")
            i += 1
            continue

        if t == "answer_start":
            # Bold only the YES/NO prefix, rest is regular weight
            full_text = entry["text"]
            import re as _re
            m = _re.match(r'^(YES|NO)\.?\s*', full_text)
            if m:
                prefix = m.group(0).rstrip()
                rest = full_text[m.end():]
                html.append(f"<p><strong>{safe(prefix)}</strong> {safe(rest)}</p>")
            else:
                html.append(f"<p>{para_html(entry)}</p>")
            i += 1
            continue

        if t == "body":
            li = ' class="lead-in"' if next_is_bullet(i) else ""
            html.append(f"<p{li}>{para_html(entry)}</p>")
            i += 1
            continue

        if t == "bullet":
            if not in_list:
                html.append("<ul>")
                in_list = True
            html.append(f"<li>{para_html(entry)}</li>")
            i += 1
            continue

        if t in ("figure_caption", "table_caption"):
            html.append(f'<div class="caption">{safe(text)}</div>')
            i += 1
            continue

        # Default
        html.append(f"<p>{para_html(entry)}</p>")
        i += 1

    if in_list:
        html.append("</ul>")

    # Flush any remaining images after the last section entry
    if sections:
        last_idx = sections[-1].get("index", -1)
        insert_pending_images(last_idx, None, html)

    # ── Q&A ────────────────────────────────────────────────
    if data.get("qa_pairs"):
        html.append("<hr>")
        for qa in data["qa_pairs"]:
            q_entry = qa["question"]
            q_prefix = ""
            if q_entry.get("numId"):
                nid = q_entry["numId"]
                num_counters[nid] = num_counters.get(nid, 0) + 1
                q_prefix = f"{num_counters[nid]}. "
            html.append(f'<div class="qa-question">{safe(q_prefix)}{para_html(q_entry)}</div>')
            for ans in qa["answer_parts"]:
                html.append(f'<div class="qa-answer">{para_html(ans)}</div>')

    # ── References ─────────────────────────────────────────
    if data.get("references"):
        html.append('</div><div class="page">')
        html.append('<div class="ref-heading">References</div>')
        for ref in data["references"]:
            html.append(f'<div class="reference">{safe(ref["text"])}</div>')

    html.append("</div>")  # close last page
    html.append("</body>")
    html.append("</html>")

    # ── Inject footer into every content page ─────────────
    # Build full title: "Cancer-Resistance and Interventions to Mitigate Resistance"
    footer_title = "-".join(title_parts)
    if subtitle:
        footer_title += " " + subtitle
    # Date from cover or fallback
    clean_date = date_str.replace("Updated ", "").strip() if date_str else ""
    if not clean_date:
        # Try to extract from filename or use cover date
        clean_date = "03/04/2026"
    footer_html = (
        f'<div class="page-footer">'
        f'<span class="footer-title">{safe(footer_title)}</span>'
        f' <span class="footer-date">({safe(clean_date)})</span>'
        f'</div>'
    )

    # Insert footer before every closing </div> of a content page
    # Pages are: <div class="page">...</div> (not cover-page)
    output = "\n".join(html)
    # Replace page close tags with footer + close, but skip cover page
    output = output.replace(
        '</div><div class="page">',
        f'{footer_html}</div><div class="page">'
    )
    # Add footer to the last page (before final </div></body>)
    output = output.replace(
        '</div>\n</body>',
        f'{footer_html}</div>\n</body>'
    )
    with open(str(out_path), "w", encoding="utf-8") as f:
        f.write(output)
    print(f"Done → {out_path}")


if __name__ == "__main__":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

    if len(sys.argv) < 2:
        print("Usage: python generate_html.py <path_to_docx> [--out output.html]")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    if "--out" in sys.argv:
        out_path = Path(sys.argv[sys.argv.index("--out") + 1])
    else:
        out_path = docx_path.with_suffix(".html")

    generate_html(docx_path, out_path)

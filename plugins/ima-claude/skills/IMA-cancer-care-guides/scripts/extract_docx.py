"""
Extract structured text from a Word document for Cancer Care Guide mapping.

Usage:
    /c/Python313/python.exe extract_docx.py <path_to_docx> [--json]

Outputs a structured representation of the document with:
- Heading hierarchy (H1, H2, H3)
- Body paragraphs with bold run markers
- Reference citations
- Q&A detection
- Word comments (if any @SLOT markers exist)

Python path: /c/Python313/python.exe
Required: pip install python-docx
"""

import sys
import json
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def extract_comments(doc):
    """Extract Word comments and their anchored paragraph indices."""
    comments = {}
    # Comments are stored in the comments part of the docx
    try:
        comments_part = doc.part.package.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        )
        if comments_part is None:
            return comments

        from lxml import etree
        tree = etree.fromstring(comments_part.blob)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        for comment in tree.findall('.//w:comment', ns):
            comment_id = comment.get(qn('w:id'))
            text_parts = []
            for p in comment.findall('.//w:t', ns):
                if p.text:
                    text_parts.append(p.text)
            text = ''.join(text_parts).strip()
            if text.startswith('@SLOT:'):
                slot_name = text.replace('@SLOT:', '').strip()
                comments[comment_id] = slot_name
    except Exception:
        pass  # No comments part or parsing error

    return comments


def get_comment_refs(paragraph):
    """Get comment IDs referenced by a paragraph (via commentRangeStart)."""
    refs = []
    for elem in paragraph._element.iter():
        if elem.tag.endswith('commentRangeStart'):
            comment_id = elem.get(qn('w:id'))
            if comment_id:
                refs.append(comment_id)
    return refs


def extract_runs_with_formatting(paragraph):
    """Extract text runs with bold/italic markers."""
    runs = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        runs.append({
            'text': text,
            'bold': run.bold or False,
            'italic': run.italic or False,
        })
    return runs


def get_num_info(paragraph):
    """Extract Word numbering info (numId, ilvl) if present."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        return None, None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(qn('w:ilvl'))
    numId_el = numPr.find(qn('w:numId'))
    ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'
    numId = numId_el.get(qn('w:val')) if numId_el is not None else None
    return numId, int(ilvl)


def has_page_break(para):
    """Check if a paragraph contains an explicit page break (w:br type=page)."""
    for run in para._element.findall(qn('w:r')):
        for br in run.findall(qn('w:br')):
            if br.get(qn('w:type'), '') == 'page':
                return True
    return False


def is_all_bold(para):
    """Check if all non-empty runs in a paragraph are bold."""
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs:
        return False
    return all(r.bold for r in text_runs)


def classify_paragraph(para, runs_info):
    """Classify a paragraph's role in the document."""
    text = para.text.strip()
    style_name = para.style.name if para.style else ''

    if not text or not text.strip('\u200b\u200c\u200d\ufeff\u00a0'):
        return 'empty', None

    # Heading detection — Word heading styles
    if style_name.startswith('Heading'):
        try:
            level = int(style_name.replace('Heading ', '').strip())
        except ValueError:
            level = 1
        return f'h{level}', None

    # Title style
    if style_name == 'Title':
        return 'h1', None

    # Date detection — must come before bold heading check since dates can be bold
    if re.match(r'^Updated\s+\w+\s+(\d{1,2}\w{0,2},?\s+)?\d{4}', text):
        return 'date', None

    # YES/NO answer lines — must come before bold heading check
    # These are bold answer lines under numbered Q&A headings
    if re.match(r'^(YES|NO)[.\s]', text):
        return 'answer_start', text[:3].rstrip('.')

    # Bold body text = heading (common in IMA medical docs)
    # Short all-bold Body/Default paragraphs are treated as section headings
    if is_all_bold(para) and len(text) < 120:
        # First bold paragraph is likely the title
        # "Table N." lines are table captions, not headings
        if re.match(r'^Table\s+\d+', text):
            return 'table_caption', None
        return 'heading_bold', None

    # Q&A detection
    if re.match(r'^Q[:\.]?\s', text) or (runs_info and runs_info[0].get('bold') and '?' in text):
        return 'question', None

    # Reference detection
    if re.match(r'^\d+\.\s+\w', text):
        return 'reference', None

    # Figure caption
    if re.match(r'^Figure\s+\d+', text):
        return 'figure_caption', None

    # Table caption (bold + starts with "Table N")
    if re.match(r'^Table\s+\d+', text):
        return 'table_caption', None

    # Disclaimer detection
    if 'complementary approach' in text.lower() or 'not intended as a comprehensive' in text.lower():
        return 'disclaimer', None

    # Warning box detection
    if 'patients should not treat themselves' in text.lower():
        return 'warning', None

    # Author detection
    if re.search(r'M\.?D\.?', text) and ('Marik' in text or 'Hope' in text or len(text) < 100):
        return 'author', None

    # Bullet/list paragraph detection — Word List Bullet style or numPr XML
    style_lower = style_name.lower()
    if 'list bullet' in style_lower or 'list number' in style_lower:
        return 'bullet', None
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:numPr')) is not None:
        return 'bullet', None

    return 'body', None


def extract_document(docx_path):
    """Extract structured content from a Word document."""
    doc = Document(docx_path)
    comments = extract_comments(doc)

    sections = []
    current_section = None
    references = []
    in_references = False
    qa_pairs = []
    current_qa = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Detect explicit page breaks (w:br type=page) on empty paragraphs
        if not text:
            if has_page_break(para):
                sections.append({
                    'index': i, 'type': 'page_break', 'text': '',
                    'runs': [], 'style': ''
                })
            continue

        runs = extract_runs_with_formatting(para)
        para_type, meta = classify_paragraph(para, runs)

        # Check for @SLOT comments
        comment_refs = get_comment_refs(para)
        slot_override = None
        for ref in comment_refs:
            if ref in comments:
                slot_override = comments[ref]
                break

        # Extract numbering info
        numId, ilvl = get_num_info(para)

        # Build paragraph entry
        entry = {
            'index': i,
            'type': para_type,
            'text': text,
            'runs': runs,
            'style': para.style.name if para.style else '',
        }
        if has_page_break(para):
            entry['page_break'] = True
        if numId is not None:
            entry['numId'] = numId
            entry['ilvl'] = ilvl
        if slot_override:
            entry['slot_override'] = slot_override
        if meta:
            entry['meta'] = meta

        # Track references section
        if text.lower() == 'references':
            in_references = True
            entry['type'] = 'ref_heading'
            sections.append(entry)
            continue

        if in_references and para_type == 'reference':
            references.append(entry)
            continue
        elif in_references and para_type != 'reference' and para_type != 'empty':
            # Might still be a reference without the number prefix
            if references:  # We already have some references
                references.append(entry)
                continue
            in_references = False

        # Track Q&A pairs
        if para_type == 'question':
            if current_qa:
                qa_pairs.append(current_qa)
            current_qa = {'question': entry, 'answer_parts': []}
            continue
        elif current_qa and (para_type in ('answer_start', 'body')):
            current_qa['answer_parts'].append(entry)
            continue
        elif current_qa and para_type == 'question':
            qa_pairs.append(current_qa)
            current_qa = {'question': entry, 'answer_parts': []}
            continue
        elif current_qa and para_type.startswith('h'):
            qa_pairs.append(current_qa)
            current_qa = None

        sections.append(entry)

    # Flush remaining Q&A
    if current_qa:
        qa_pairs.append(current_qa)

    return {
        'source_file': str(docx_path),
        'total_paragraphs': len(doc.paragraphs),
        'sections': sections,
        'qa_pairs': qa_pairs,
        'references': references,
        'slot_overrides': {v: k for k, v in comments.items()} if comments else {},
    }


def print_readable(data):
    """Print human-readable summary."""
    print(f"Source: {data['source_file']}")
    print(f"Total paragraphs: {data['total_paragraphs']}")
    print(f"Sections: {len(data['sections'])}")
    print(f"Q&A pairs: {len(data['qa_pairs'])}")
    print(f"References: {len(data['references'])}")
    print(f"@SLOT overrides: {len(data['slot_overrides'])}")
    print()

    print("=== DOCUMENT STRUCTURE ===")
    for entry in data['sections']:
        prefix = ''
        if entry['type'] == 'h1':
            prefix = '# '
        elif entry['type'] == 'h2':
            prefix = '## '
        elif entry['type'] == 'h3':
            prefix = '### '
        elif entry['type'] == 'heading_bold':
            prefix = '** '
        elif entry['type'] == 'disclaimer':
            prefix = '[DISCLAIMER] '
        elif entry['type'] == 'warning':
            prefix = '[WARNING] '
        elif entry['type'] == 'date':
            prefix = '[DATE] '
        elif entry['type'] == 'author':
            prefix = '[AUTHOR] '
        elif entry['type'] == 'figure_caption':
            prefix = '[FIGURE] '
        elif entry['type'] == 'ref_heading':
            prefix = '[REFERENCES] '

        slot = entry.get('slot_override', '')
        slot_str = f' → @SLOT:{slot}' if slot else ''

        text_preview = entry['text'][:120] + ('...' if len(entry['text']) > 120 else '')
        print(f"  {prefix}{text_preview}{slot_str}")

    if data['qa_pairs']:
        print()
        print("=== Q&A PAIRS ===")
        for qa in data['qa_pairs']:
            q_text = qa['question']['text'][:100]
            a_count = len(qa['answer_parts'])
            print(f"  Q: {q_text}")
            print(f"     ({a_count} answer paragraph(s))")

    if data['references']:
        print()
        print(f"=== REFERENCES ({len(data['references'])} total) ===")
        for ref in data['references'][:3]:
            print(f"  {ref['text'][:100]}...")
        if len(data['references']) > 3:
            print(f"  ... and {len(data['references']) - 3} more")


if __name__ == '__main__':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <path_to_docx> [--json]")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    output_json = '--json' in sys.argv

    data = extract_document(docx_path)

    if output_json:
        # JSON output for Claude to consume
        print(json.dumps(data, indent=2, ensure_ascii=False))
    else:
        # Human-readable output
        print_readable(data)
